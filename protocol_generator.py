"""
protocol_generator.py — генератор протокола .docx из MD-задач собрания ЗПР

Использует корпоративный бланк: _шаблоны/Протокол_бланк.docx
Создаёт копию бланка и заполняет её данными из задач.

Структура бланка:
  Heading:  "Протокол рабочего собрания DD.MM.YYYY г."
  Table 0:  Общая информация (объект, предмет рассмотрения)
  Table 1:  Участники (№ / Организация / Представитель)
  Table 2:  Обсудили — Выполнено (выполненные задачи)
  Table 3:  Обсудили — Надо подготовить (открытые задачи)

Запуск:
  python protocol_generator.py "ПОДРЯДЧИКИ/Бюро82/Собрания/2026-04-17 Рабочее собрание"
"""

import sys
import re
import copy
import argparse
from datetime import datetime
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

from config import BASE_DIR

TEMPLATE_PATH = BASE_DIR / "_шаблоны" / "Протокол_бланк.docx"

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Установите python-docx: pip install python-docx")
    sys.exit(1)


# ─── Чтение MD-файлов ─────────────────────────────────────────────────────────

def parse_frontmatter(text: str) -> dict:
    m = re.match(r"^---\n(.*?)\n---", text, re.DOTALL)
    if not m:
        return {}
    fm = {}
    for line in m.group(1).splitlines():
        if ":" in line:
            key, _, val = line.partition(":")
            fm[key.strip()] = val.strip().strip('"')
    return fm


def parse_participants_file(path: Path) -> dict:
    text = path.read_text(encoding="utf-8")
    fm = parse_frontmatter(text)
    rows = []
    for line in text.splitlines():
        line = line.strip()
        if line.startswith("|") and "---" not in line:
            cells = [c.strip() for c in line.strip("|").split("|")]
            if len(cells) >= 4 and cells[0] not in ("№", ""):
                # Очищаем <br> и лишние пробелы из всех полей
                clean = lambda s: re.sub(r"<br\s*/?>", "", s).strip()
                rows.append({
                    "num":  clean(cells[0]),
                    "org":  clean(cells[1]),
                    "role": clean(cells[2]),
                    "name": clean(cells[3]),
                })
    return {
        "meeting": fm.get("meeting", ""),
        "date":    fm.get("date", ""),
        "project": fm.get("project", ""),
        "rows":    rows,
    }


def load_tasks(tasks_dir: Path) -> list[dict]:
    tasks = []
    for f in sorted(tasks_dir.glob("ПРОТ-*.md")):
        text = f.read_text(encoding="utf-8")
        fm = parse_frontmatter(text)
        # explanation из тела если в frontmatter пустое
        body = re.search(r"^---\n.*?\n---\n\s*#[^\n]*\n(.+?)(?:\n## |$)", text, re.DOTALL)
        if body and not fm.get("explanation"):
            fm["explanation"] = body.group(1).strip()
        fm["_file"] = f.name
        tasks.append(fm)
    return tasks


# ─── Утилиты работы с таблицей ────────────────────────────────────────────────

def clear_table_data_rows(table, keep_rows: int = 1):
    """Удаляет все строки кроме первых keep_rows (заголовок)."""
    while len(table.rows) > keep_rows:
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)


def copy_row_format(table, source_row_idx: int = 1):
    """Возвращает XML-элемент строки-образца для копирования форматирования."""
    if len(table.rows) > source_row_idx:
        return table.rows[source_row_idx]._tr
    return table.rows[-1]._tr


def unique_cells(row):
    """Возвращает только уникальные ячейки строки (merged cells не дублируются)."""
    seen = set()
    result = []
    for cell in row.cells:
        tc_id = id(cell._tc)
        if tc_id not in seen:
            seen.add(tc_id)
            result.append(cell)
    return result


def write_cell(cell, text: str):
    """Записывает текст в ячейку, сохраняя форматирование первого рана."""
    for p in cell.paragraphs:
        # Очищаем все раны
        for run in p.runs:
            run.text = ""
        # Пишем в первый ран (сохраняет шрифт/размер из шаблона)
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)
        # Убираем лишние параграфы (в merged cell может быть несколько)
        break


def add_table_row(table, values: list[str], template_tr=None) -> None:
    """Добавляет строку с копированием форматирования из template_tr."""
    if template_tr is not None:
        new_tr = copy.deepcopy(template_tr)
        table._tbl.append(new_tr)
        row = table.rows[-1]
        cells = unique_cells(row)
        for i, val in enumerate(values):
            if i < len(cells):
                write_cell(cells[i], val)
    else:
        row = table.add_row()
        cells = unique_cells(row)
        for i, val in enumerate(values):
            if i < len(cells):
                cells[i].text = val


def set_cell_text(cell, text: str):
    """Устанавливает текст ячейки, сохраняя форматирование первого рана."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = text
            return
        p.add_run(text)
        return
    cell.paragraphs[0].add_run(text)


def format_date(date_str: str) -> str:
    """YYYY-MM-DD → DD.MM.YYYY"""
    m = re.match(r"(\d{4})-(\d{2})-(\d{2})", date_str or "")
    if m:
        return f"{m.group(3)}.{m.group(2)}.{m.group(1)}"
    return date_str or "—"


# ─── Построение документа ─────────────────────────────────────────────────────

def build_protocol(info: dict, tasks: list[dict], output_path: Path):
    if not TEMPLATE_PATH.exists():
        print(f"❌ Бланк не найден: {TEMPLATE_PATH}")
        sys.exit(1)

    doc = Document(str(TEMPLATE_PATH))

    date_display = format_date(info["date"])
    open_tasks = [t for t in tasks if t.get("status", "open") != "done"]
    done_tasks  = [t for t in tasks if t.get("status", "open") == "done"]

    # ── Заголовок (первый непустой параграф Heading 2) ──
    for p in doc.paragraphs:
        if p.style.name == "Heading 2" and p.text.strip():
            # Собираем весь текст в первый ран, остальные очищаем
            new_text = f"Протокол рабочего собрания {date_display} г."
            if p.runs:
                p.runs[0].text = new_text
                for run in p.runs[1:]:
                    run.text = ""
            else:
                p.add_run(new_text)
            break

    tables = doc.tables

    # ── Таблица 0: Общая информация ──
    if len(tables) > 0:
        t0 = tables[0]
        if len(t0.rows) > 1:
            write_cell(unique_cells(t0.rows[1])[2], info.get("project", ""))
        if len(t0.rows) > 2:
            write_cell(unique_cells(t0.rows[2])[2], info.get("meeting", ""))

    # ── Таблица 1: Участники ──
    if len(tables) > 1:
        t1 = tables[1]
        tmpl_tr = copy_row_format(t1, 1) if len(t1.rows) > 1 else None
        clear_table_data_rows(t1, keep_rows=1)
        for prow in info["rows"]:
            name = prow["name"].strip()
            role = prow["role"].strip()
            rep  = f"{name} — {role}" if name and role else (name or role)
            add_table_row(t1, [prow["num"], prow["org"], rep], template_tr=tmpl_tr)

    # ── Таблица 2: Выполнено ──
    if len(tables) > 2:
        t2 = tables[2]
        tmpl_tr = copy_row_format(t2, 1) if len(t2.rows) > 1 else None
        clear_table_data_rows(t2, keep_rows=1)
        if done_tasks:
            for i, task in enumerate(done_tasks, 1):
                add_table_row(t2,
                    [str(i), task.get("title", ""), task.get("explanation", ""), "✅"],
                    template_tr=tmpl_tr)
        else:
            add_table_row(t2, ["—", "—", "Нет выполненных пунктов", "—"], template_tr=tmpl_tr)

    # ── Таблица 3: Надо подготовить ──
    if len(tables) > 3:
        t3 = tables[3]
        tmpl_tr = copy_row_format(t3, 1) if len(t3.rows) > 1 else None
        clear_table_data_rows(t3, keep_rows=1)
        for i, task in enumerate(open_tasks, 1):
            num   = f"{4}.{i}"
            due   = format_date(task.get("due", "")) if task.get("due") else "—"
            add_table_row(t3,
                [num, task.get("title", ""), task.get("explanation", ""),
                 task.get("assignee", "—"), due],
                template_tr=tmpl_tr)

    # ── Дата создания ──
    today = datetime.today().strftime("%Y-%m-%d")
    for p in doc.paragraphs:
        if "Создано:" in p.text:
            if p.runs:
                p.runs[0].text = f"Создано: {today}"
                for r in p.runs[1:]:
                    r.text = ""
            break

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ─── main ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Генератор протокола .docx из MD-задач")
    parser.add_argument("meeting_dir", help="Путь к папке собрания")
    args = parser.parse_args()

    meeting_folder = Path(args.meeting_dir)
    if not meeting_folder.is_absolute():
        meeting_folder = BASE_DIR / meeting_folder
    if not meeting_folder.exists():
        print(f"❌ Папка не найдена: {meeting_folder}")
        sys.exit(1)

    participants_file = meeting_folder / "Участники.md"
    if not participants_file.exists():
        print("❌ Файл Участники.md не найден")
        sys.exit(1)

    tasks_dir = meeting_folder / "Задачи"
    if not tasks_dir.exists() or not list(tasks_dir.glob("ПРОТ-*.md")):
        print("❌ Папка Задачи/ пуста или не существует")
        sys.exit(1)

    info  = parse_participants_file(participants_file)
    tasks = load_tasks(tasks_dir)

    open_tasks = [t for t in tasks if t.get("status", "open") != "done"]
    done_tasks  = [t for t in tasks if t.get("status", "open") == "done"]

    print(f"📋 Собрание:   {info['meeting']}")
    print(f"📅 Дата:       {info['date']}")
    print(f"👥 Участников: {len(info['rows'])}")
    print(f"✅ Выполнено:  {len(done_tasks)}")
    print(f"📌 Открытых:   {len(open_tasks)}")

    # Папка Протоколы/ — два уровня вверх от папки собрания
    protokoly_dir = meeting_folder.parent.parent / "Протоколы"

    date_str = info["date"]
    m = re.match(r"(\d{2})\.(\d{2})\.(\d{4})", date_str)
    date_iso = f"{m.group(3)}-{m.group(2)}-{m.group(1)}" if m else date_str
    output_path = protokoly_dir / f"Протокол_{date_iso}.docx"

    print(f"\n💾 Сохраняем: {output_path.relative_to(BASE_DIR)}")
    build_protocol(info, tasks, output_path)
    print(f"✅ Готово: {output_path.name}")


if __name__ == "__main__":
    main()
